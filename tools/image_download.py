"""飞书文件下载工具 — 通过 file_key 下载文件内容到本地临时路径。"""
import os
import tempfile
import urllib.request
from core.context import ToolManifest

MANIFEST = ToolManifest(name="feishu_file_download", description="从飞书下载文件到临时目录")


def download_file(token: str, message_id: str, file_key: str, file_name: str) -> str:
    """
    下载飞书消息中的文件，返回本地临时文件路径。
    失败时抛出异常。
    """
    url = f"https://open.feishu.cn/open-apis/im/v1/messages/{message_id}/resources/{file_key}?type=file"
    req = urllib.request.Request(url, headers={
        "Authorization": f"Bearer {token}",
        "User-Agent": "Mozilla/5.0",
    })
    with urllib.request.urlopen(req, timeout=60) as resp:
        data = resp.read()

    suffix = os.path.splitext(file_name)[1] or '.bin'
    tmp = tempfile.NamedTemporaryFile(delete=False, suffix=suffix)
    tmp.write(data)
    tmp.close()
    return tmp.name


def extract_text_from_file(file_path: str) -> str:
    """
    从本地文件提取纯文本。
    支持：.docx（Word）/ .pdf / .txt / .md
    失败时返回空字符串。
    """
    ext = os.path.splitext(file_path)[1].lower()

    if ext == '.docx':
        return _extract_docx(file_path)
    elif ext == '.pdf':
        return _extract_pdf(file_path)
    elif ext in ('.txt', '.md'):
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            return f.read()
    else:
        return ''


def _extract_docx(path: str) -> str:
    import docx
    doc = docx.Document(path)
    parts = []
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text.strip())
    # 提取表格内容
    for table in doc.tables:
        for row in table.rows:
            row_text = '  |  '.join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                parts.append(row_text)
    return '\n'.join(parts)


def _extract_pdf(path: str) -> str:
    import pdfplumber
    pages = []
    with pdfplumber.open(path) as pdf:
        for page in pdf.pages:
            text = page.extract_text()
            if text:
                pages.append(text.strip())
    return '\n\n'.join(pages)


def handle(ctx):
    return ctx

def download_image(token: str, message_id: str, image_key: str) -> bytes:
    """
    下载飞书消息中的图片，返回字节数据。
    """
    url = f"https://open.feishu.cn/open-apis/im/v1/messages/{message_id}/resources/{image_key}?type=image"
    req = urllib.request.Request(url, headers={
        "Authorization": f"Bearer {token}",
        "User-Agent": "Mozilla/5.0",
    })
    with urllib.request.urlopen(req, timeout=60) as resp:
        return resp.read()
